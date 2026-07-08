from __future__ import annotations

from io import BytesIO
import zipfile

from docx import Document
from docx.shared import Inches
from PIL import Image

from services.docx_exporter import export_docx
from services.docx_ir import parse_docx_ir, render_docx_ir
from services.docx_parser import extract_segments


def _make_docx_with_inline_image() -> bytes:
    image_buffer = BytesIO()
    Image.new("RGB", (24, 24), color=(255, 0, 0)).save(image_buffer, format="PNG")
    image_buffer.seek(0)

    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Step 1 ")
    paragraph.add_run().add_picture(image_buffer, width=Inches(0.25))
    paragraph.add_run(" install bracket")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _make_docx_with_multiline_table_cell() -> bytes:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.add_run("Line 1")
    paragraph.add_run().add_break()
    paragraph.add_run("Line 2")
    paragraph.add_run().add_break()
    paragraph.add_run("Line 3")

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _document_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        return archive.read("word/document.xml").decode("utf-8", errors="ignore")


def test_export_docx_preserves_inline_drawings_when_replacing_text() -> None:
    original = _make_docx_with_inline_image()
    segments = extract_segments(original)
    assert len(segments) == 1

    translated = export_docx(
        original,
        segments,
        [{"translation": "Step One install bracket translated"}],
    )

    original_xml = _document_xml(original)
    translated_xml = _document_xml(translated)

    assert original_xml.count("<w:drawing") == 1
    assert translated_xml.count("<w:drawing") == 1
    assert "Step One install bracket translated" in translated_xml


def test_export_docx_sets_run_language_to_target_lang() -> None:
    doc = Document()
    doc.add_paragraph("安装支架")
    output = BytesIO()
    doc.save(output)

    original = output.getvalue()
    segments = extract_segments(original)
    translated = export_docx(
        original,
        segments,
        [{"translation": "Install bracket"}],
        target_lang="en-US",
    )

    translated_xml = _document_xml(translated)
    assert 'w:lang w:val="en-US"' in translated_xml


def test_parse_docx_ir_keeps_docx_segments_compatible_with_existing_parser() -> None:
    original = _make_docx_with_inline_image()

    ir = parse_docx_ir(original)
    segments = extract_segments(original)

    assert ir["kind"] == "docx"
    assert ir["segments"] == segments
    assert ir["nodes"][0]["hasDrawing"] is True


def test_render_docx_ir_matches_existing_export_path() -> None:
    original = _make_docx_with_inline_image()
    ir = parse_docx_ir(original)
    segments = extract_segments(original)
    translated_segments = [{"translation": "Step One install bracket translated"}]

    expected = export_docx(original, segments, translated_segments)
    rendered = render_docx_ir(original, ir, translated_segments)

    assert _document_xml(rendered) == _document_xml(expected)


def test_parse_docx_ir_splits_multiline_table_cells_into_line_segments() -> None:
    original = _make_docx_with_multiline_table_cell()

    ir = parse_docx_ir(original)
    segments = extract_segments(original)

    assert len(ir["segments"]) == 3
    assert len(segments) == 3
    assert [segment["plain_text"] for segment in ir["segments"]] == ["Line 1", "Line 2", "Line 3"]
    assert [segment["line_index"] for segment in ir["segments"]] == [0, 1, 2]
    assert [segment["line_count"] for segment in ir["segments"]] == [3, 3, 3]


def test_render_docx_ir_rebuilds_multiline_table_cells_in_order() -> None:
    original = _make_docx_with_multiline_table_cell()
    ir = parse_docx_ir(original)
    translated_segments = [
        {"translation": "First line"},
        {"translation": "Second line"},
        {"translation": "Third line"},
    ]

    rendered = render_docx_ir(original, ir, translated_segments)
    translated_xml = _document_xml(rendered)

    assert "First line" in translated_xml
    assert "Second line" in translated_xml
    assert "Third line" in translated_xml
    assert translated_xml.index("First line") < translated_xml.index("Second line") < translated_xml.index("Third line")
